"""Report generator that produces JSON and HTML outputs."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from jinja2 import Environment
from jinja2 import FileSystemLoader
from jinja2 import select_autoescape

TEMPLATE_DIR = Path(__file__).parent / "templates"


env = Environment(
    loader=FileSystemLoader(str(TEMPLATE_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)


class ReportGenerator:
    """Generates analysis reports in multiple formats."""
    SUPPORTED_FORMATS = {"json", "html", "md", "pdf", "docx"}

    def _render_html(self, data: dict[str, Any]) -> str:
        template = env.get_template("report.html")
        return template.render(results=data)

    # ---------------------------- public API ---------------------------- #
    def generate(self, base_path: Path, results: dict[str, Any], formats: list[str] | None = None) -> list[Path]:
        """Generate report files.

        Parameters
        ----------
        base_path : Path
            Path provided by user *without* extension. We will add extension per format.
        results : dict
            Analysis results.
        formats : list[str] | None
            Allowed strings: ``json``, ``html``, ``md``, ``pdf``, ``docx``. If ``None`` we default to
            ``["json", "html"]``.

        Returns
        -------
        list[Path]
            Paths of written report files.
        """

        if not formats:
            formats = ["json", "html"]

        # Validate that all requested formats are supported
        unsupported = set(formats) - self.SUPPORTED_FORMATS
        if unsupported:
            raise ValueError(
                f"Unsupported report format(s): {', '.join(unsupported)}. "
                f"Supported formats are: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        written: list[Path] = []

        # JSON ---------------------------------------------------------
        if "json" in formats:
            json_path = base_path.with_suffix(".json")
            json_path.write_text(json.dumps(results, indent=2), encoding="utf-8")
            written.append(json_path)

        # HTML ---------------------------------------------------------
        html_content = self._render_html(results)
        if "html" in formats:
            html_path = base_path.with_suffix(".html")
            html_path.write_text(html_content, encoding="utf-8")
            written.append(html_path)

        # Markdown -----------------------------------------------------
        if "md" in formats:
            md_path = base_path.with_suffix(".md")
            md_parts: list[str] = ["# Codebase Analysis Report\n"]

            # ------------- Human-readable sections -------------
            parser = results.get("parser_results", {})
            md_parts.append("## Parser Results")
            if parser:
                md_parts.append(
                    f"- **Files analyzed:** {parser.get('file_count', 0)}\n"
                    f"- **Functions detected:** {parser.get('function_count', 0)}\n"
                    f"- **Call-graph nodes:** {parser.get('call_graph_nodes', 0)}\n"
                    f"- **Call-graph edges:** {parser.get('call_graph_edges', 0)}\n"
                )
                summary = parser.get("gemini_summary", "")
                if summary:
                    md_parts.append("- **Gemini summary:**\n")
                    for line in summary.strip().splitlines():
                        md_parts.append(f"  > {line}")
            else:
                md_parts.append("No parser results available.\n")

            perf = results.get("performance_issues", {})
            md_parts.append("\n## Performance Issues")
            if perf and perf.get("count", 0):
                for issue in perf.get("issues", []):
                    md_parts.append(
                        f"- {issue.get('file', '')}: {issue.get('detail', issue)}"
                    )
            else:
                md_parts.append("No significant performance issues detected.\n")

            sec = results.get("security_findings", {})
            md_parts.append("\n## Security Findings")
            if sec and sec.get("count", 0):
                for finding in sec.get("issues", []):
                    md_parts.append(
                        f"- {finding.get('file', '')}: {finding.get('issue', '')} – {finding.get('detail', '')}"
                    )
            else:
                md_parts.append("No critical security findings detected.\n")

            md_path.write_text("\n".join(md_parts), encoding="utf-8")
            written.append(md_path)

        # PDF ----------------------------------------------------------
        if "pdf" in formats:
            try:
                from fpdf import FPDF  # type: ignore

                pdf_path = base_path.with_suffix(".pdf")

                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Courier", size=8)

                # Add each line of the Markdown-ish report to the PDF
                for line in self._render_markdown_like(results).splitlines():
                    for chunk in self._wrap_line(line, width=80):
                        pdf.multi_cell(0, 5, text=chunk)

                pdf.output(str(pdf_path))
                written.append(pdf_path)
            except Exception as exc:  # pragma: no cover
                import logging

                logging.getLogger(__name__).warning(
                    "Skipping PDF export: %s (PDF generation failed)", exc,
                )

        # DOCX ---------------------------------------------------------
        if "docx" in formats:
            try:
                from docx import Document  # type: ignore

                document = Document()
                document.add_heading("Codebase Analysis Report", level=1)
                for section, data in results.items():
                    document.add_heading(section.replace("_", " ").title(), level=2)
                    document.add_paragraph(json.dumps(data, indent=2))

                docx_path = base_path.with_suffix(".docx")
                document.save(str(docx_path))
                written.append(docx_path)
            except Exception as exc:  # pragma: no cover
                import logging

                logging.getLogger(__name__).warning(
                    "Skipping DOCX export: %s. Install python-docx to enable.", exc,
                )

        return written

    # -----------------------------------------------------------------
    def _render_markdown_like(self, results: dict[str, Any]) -> str:
        """Return a simple Markdown-style text for PDF fallback."""
        md_parts: list[str] = ["# Codebase Analysis Report\n"]
        for section, data in results.items():
            md_parts.append(f"## {section.replace('_', ' ').title()}")
            md_parts.append(json.dumps(data, indent=2))
            md_parts.append("")
        return "\n".join(md_parts)

    # ------------------------ helpers ----------------------------- #
    @staticmethod
    def _wrap_line(line: str, width: int = 80) -> list[str]:
        """Return a list of soft-wrapped chunks no longer than *width* chars."""
        if len(line) <= width:
            return [line]
        chunks: list[str] = []
        start = 0
        while start < len(line):
            chunks.append(line[start : start + width])
            start += width
        return chunks
